import os
import fitz  
import pdfplumber
import docx
from typing import List, Dict
from langchain_core.documents import Document
import pandas as pd
import numpy as np
class SimpleParser:
    def __init__(self, file_path: str, debug: bool = False):
        self.file_path = file_path
        self.debug = debug
        self.extension = os.path.splitext(file_path)[1].lower()
        print(f"Initializing SimpleParser with file: {file_path}, extension: {self.extension}")

        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if self.extension not in [".pdf", ".docx", ".csv", ".xlsx"]:
            raise ValueError("Unsupported file format. Only .pdf, .docx, .csv, and .xlsx are supported.")


    def extract_text(self) -> List[Document]:
        if self.extension == ".pdf":
            return self._extract_text_from_pdf()
        elif self.extension == ".docx":
            return self._extract_text_from_docx()
        elif self.extension in [".csv", ".xlsx"]:
            return self._extract_from_spreadsheet()
        else:
            raise ValueError(f"Unsupported extension: {self.extension}")


    def extract_tables(self) -> List[Document]:
        if self.extension == ".pdf":
            return self._extract_tables_from_pdf()
        elif self.extension == ".docx":
            return self._extract_tables_from_docx()

    def _extract_text_from_pdf(self) -> List[Document]:
        doc = fitz.open(self.file_path)
        documents = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"page": i + 1, "type": "text", "source": self.file_path}
                ))
        doc.close()
        return documents

    def _extract_tables_from_pdf(self) -> List[Document]:
        tables = []
        with pdfplumber.open(self.file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        # Convert table to readable string
                        table_str = "\n".join(["\t".join(row) for row in table if row])
                        tables.append(Document(
                            page_content=table_str,
                            metadata={"page": i + 1, "type": "table", "source": self.file_path}
                        ))
        return tables

    def _extract_text_from_docx(self) -> List[Document]:
        doc = docx.Document(self.file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [
            Document(
                page_content=full_text,
                metadata={"type": "text", "source": self.file_path}
            )
        ]

    def _extract_tables_from_docx(self) -> List[Document]:
        doc = docx.Document(self.file_path)
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            table_str = "\n".join(["\t".join(row) for row in table_data])
            tables.append(Document(
                page_content=table_str,
                metadata={"type": "table", "table_index": i, "source": self.file_path}
            ))
        return tables
    


    def _extract_from_spreadsheet(self) -> List[Document]:
        """
        Trích xuất dữ liệu từ file CSV/XLSX và trả về dưới dạng list các Document.
        Mỗi hàng trong bảng sẽ thành một Document riêng.
        """
        if self.extension == ".csv":
            df = pd.read_csv(self.file_path)
            file_type = "csv"
        elif self.extension == ".xlsx":
            df = pd.read_excel(self.file_path)
            file_type = "excel"
        else:
            raise ValueError("Unsupported spreadsheet format.")

        rows = df.to_numpy()
        row_indices = np.arange(len(rows))
        sources = np.repeat(self.file_path, len(rows))

        row_texts = np.apply_along_axis(
            lambda x: "-".join(x.astype(str)), 
            axis=1,
            arr=rows
        )

        documents = [
            Document(
                page_content=text,
                metadata={
                    "type": file_type,
                    "row_index": int(idx),
                    "source": source
                }
            ) for text, idx, source in zip(row_texts, row_indices, sources)
        ]

        return documents